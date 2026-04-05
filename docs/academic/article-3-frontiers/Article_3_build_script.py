#!/usr/bin/env python3
"""Build expanded Frontiers manuscript targeting ~11,500 body words."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1); s.bottom_margin = Inches(1)
    s.left_margin = Inches(1); s.right_margin = Inches(1)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0; style.paragraph_format.space_after = Pt(0)

def h1(t):
    h = doc.add_heading(t, level=1)
    for r in h.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
def h2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
def p(t, indent=True, bold=False, italic=False):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 2.0; par.paragraph_format.space_after = Pt(6)
    if indent: par.paragraph_format.first_line_indent = Inches(0.5)
    r = par.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    r.bold=bold; r.italic=italic
def pn(t): p(t, indent=False)
def pb(): doc.add_page_break()
def fp(t): p(t, indent=False, bold=True, italic=True)
def meta(t):
    par = doc.add_paragraph()
    r = par.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(10)
    r.font.color.rgb = RGBColor(100,100,100)
def title(t):
    par = doc.add_paragraph()
    r = par.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(16); r.bold=True
def ref(t):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.15; par.paragraph_format.space_after = Pt(4)
    par.paragraph_format.first_line_indent = Inches(-0.5)
    par.paragraph_format.left_indent = Inches(0.5)
    r = par.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(11)
def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for pr in c.paragraphs:
            for r in pr.runs: r.bold=True; r.font.size=Pt(10); r.font.name='Times New Roman'
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = txt
            for pr in c.paragraphs:
                for r in pr.runs: r.font.size=Pt(10); r.font.name='Times New Roman'

# Read body text from separate file
import json
with open('/home/claude/sections.json', 'r') as f:
    S = json.load(f)

# TITLE PAGE
meta("HYPOTHESIS AND THEORY ARTICLE")
meta("Submitted to: Frontiers in Pharmacology — Neuropharmacology")
doc.add_paragraph()
title("The Terpene Pre-Loading Hypothesis: Can Cross-Botanical Inhalation of Terpenoids Prime Receptor Systems to Enhance Cannabinoid Pharmacology?")
doc.add_paragraph()
pn("Charles Pelletier-Gagné")
p("Independent Researcher, Sebastopol, California, United States", indent=False, italic=True)
doc.add_paragraph()
pn("Correspondence: Charles Pelletier-Gagné, charles@terpenetwist.com")
pn("Running title: Terpene Pre-Loading Hypothesis")
pn("Keywords: terpene, pre-loading, cannabinoid, entourage effect, adenosine A2a receptor, aromatherapy, essential oil, inhalation pharmacokinetics, CB1/CB2 heteromer, β-caryophyllene, d-limonene, linalool, myrcene, falsifiable predictions")
pn("Word count: ~11,500 (body text)")
pn("Figures/Tables: 4 Figures, 3 Tables (7 total)")
pn("Number of references: 75")
pb()

# ABSTRACT
h1("Abstract")
pn(S["abstract"])
pb()

# Build each section
for sec in S["sections"]:
    h1(sec["title"])
    for item in sec["content"]:
        if item["type"] == "h2":
            h2(item["text"])
        elif item["type"] == "p":
            p(item["text"])
        elif item["type"] == "fp":
            fp(item["text"])
        elif item["type"] == "table":
            p(item["caption"], indent=False, bold=True)
            add_table(item["headers"], item["rows"])
        elif item["type"] == "pb":
            pb()
pb()

# BACK MATTER
h1("Author Contributions")
pn("CP-G conceived the hypothesis, reviewed the literature, designed the experimental protocols, and wrote the manuscript.")
h1("Funding")
pn("This research received no external funding.")
h1("Conflict of Interest")
pn("The author declares that the research was conducted in the absence of any commercial or financial relationships that could be construed as a potential conflict of interest. The author is developing commercial applications of terpene-cannabis pairing through Terpenes, With a Twist of Aromatherapy, a consumer-facing education and product platform. This commercial interest is disclosed in the spirit of transparency but did not influence the scientific content, predictions, or experimental designs presented herein.")
h1("Acknowledgments")
pn("The author thanks the independent cannabis and terpene research community, particularly the laboratories of John Streicher (University of Arizona), Ryan Vandrey and Tory Spindle (Johns Hopkins University), and Yair Ben-Chaim (Open University of Israel), whose published work forms the experimental foundation of this hypothesis. The author also acknowledges the contributions of Ethan Russo, whose 2011 review first articulated the phytocannabinoid-terpenoid synergy vision that the pre-loading hypothesis seeks to operationalize.")
pb()

# REFERENCES
h1("References")
for r in S["references"]:
    ref(r)

doc.save("/home/claude/frontiers_v2.docx")

# Word count
body_w = 0; in_refs = False
for par in doc.paragraphs:
    t = par.text.strip()
    if t == "References": in_refs = True
    if not in_refs and t: body_w += len(t.split())
print(f"Body words: {body_w}")
print(f"Target: 11,500 | Delta: {11500-body_w:+d}")
